"""
Compliance Suite Service - Template-based, evidence-grounded document generation.

Generates five compliance documents from templates:
- Annex IV Technical Documentation
- FRIA (Article 27)
- PMM Report (Article 72)
- ISO 42001 Statement of Applicability
- AI Risk Register & CAPA Log

All paragraphs must be evidence-grounded with citations.
"""

import hashlib
import json
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from jinja2 import Environment, FileSystemLoader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models import (
    AISystem, ArtifactText, Control, Evidence, FRIA, Incident, 
    Organization, SoAItem
)


class ComplianceSuiteService:
    """Service for generating compliance documents from templates."""
    
    # Map document types to template filenames
    TEMPLATE_MAPPING = {
        "annex_iv": "annex_iv.md",
        "fria": "fria.md",
        "pmm": "pmm_report.md",
        "soa": "soa.md",
        "risk_register": "risk_register.md"
    }
    
    def __init__(self):
        """Initialize the compliance suite service."""
        # Make templates_dir relative to the project root
        project_root = Path(__file__).parent.parent.parent.parent
        self.templates_dir = project_root / settings.TEMPLATES_DIR
        self.jinja_env = Environment(
            loader=FileSystemLoader(self.templates_dir),
            autoescape=True
        )
    
    def _add_footer(self, content: str) -> str:
        """Add compliance footer to document content."""
        # Calculate SHA-256 of content
        content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
        timestamp = datetime.now(timezone.utc).isoformat()
        
        footer = f"""

---

**Prepared by Foundry AI Governance**  
SHA-256: `{content_hash}`  
Generated: {timestamp}  

"""
        return content + footer
    
    def generate_draft_documents(
        self, 
        db: Session, 
        org_id: int, 
        system_id: Optional[int] = None,
        doc_types: Optional[List[str]] = None
    ) -> Dict:
        """
        Generate draft compliance documents with evidence-grounded content.
        
        Args:
            db: Database session
            org_id: Organization ID
            system_id: Optional system ID for system-specific docs
            doc_types: List of document types to generate
            
        Returns:
            Dict with document drafts, coverage metrics, and missing sections
        """
        if doc_types is None:
            doc_types = ["annex_iv", "fria", "pmm", "soa", "risk_register"]
        
        # Get organization and system data
        org = db.query(Organization).filter(Organization.id == org_id).first()
        if not org:
            raise ValueError("Organization not found")
        
        system = None
        if system_id:
            system = db.query(AISystem).filter(
                AISystem.id == system_id, 
                AISystem.org_id == org_id
            ).first()
            if not system:
                raise ValueError("System not found")
        
        results = {"docs": []}
        
        for doc_type in doc_types:
            doc_result = self._generate_document(
                db, org, system, doc_type
            )
            results["docs"].append(doc_result)
        
        return results
    
    def _generate_document(
        self, 
        db: Session, 
        org: Organization, 
        system: Optional[AISystem], 
        doc_type: str
    ) -> Dict:
        """Generate a single document with evidence-grounded content."""
        
        # Load template using the mapping
        template_name = self.TEMPLATE_MAPPING.get(doc_type, f"{doc_type}.md")
        template = self.jinja_env.get_template(template_name)
        
        # Get evidence-grounded content
        sections = self._get_evidence_grounded_sections(
            db, org.id, system.id if system else None, doc_type
        )
        
        # Prepare template variables
        template_vars = self._prepare_template_variables(org, system, sections)
        
        # Render document
        try:
            content = template.render(**template_vars)
            # Add footer with hash and timestamp
            content = self._add_footer(content)
        except Exception as e:
            content = f"Error rendering template: {str(e)}"
        
        # Calculate coverage
        total_sections = len(sections)
        covered_sections = sum(1 for s in sections.values() if s.get("coverage", 0) > 0)
        coverage = covered_sections / total_sections if total_sections > 0 else 0
        
        # Identify missing sections
        missing = [
            key for key, section in sections.items() 
            if section.get("coverage", 0) == 0
        ]
        
        # Convert sections dict to list for Pydantic schema
        sections_list = [
            {
                "key": key,
                "coverage": section.get("coverage", 0),
                "paragraphs": section.get("paragraphs", [])
            }
            for key, section in sections.items()
        ]
        
        return {
            "type": doc_type,
            "content": content,
            "coverage": coverage,
            "sections": sections_list,
            "missing": missing
        }
    
    def _get_evidence_grounded_sections(
        self, 
        db: Session, 
        org_id: int, 
        system_id: Optional[int], 
        doc_type: str
    ) -> Dict[str, Dict]:
        """Get evidence-grounded content for document sections."""
        
        # Define section keys for each document type
        section_keys = self._get_section_keys(doc_type)
        
        sections = {}
        
        for section_key in section_keys:
            # Search for evidence in artifact_text
            evidence_snippets = self._search_evidence_snippets(
                db, org_id, system_id, section_key
            )
            
            if evidence_snippets:
                # Generate evidence-grounded paragraph
                paragraph = self._generate_evidence_paragraph(
                    evidence_snippets, section_key
                )
                sections[section_key] = {
                    "coverage": 1.0,
                    "paragraphs": [paragraph]
                }
            else:
                # No evidence found - mark as missing
                sections[section_key] = {
                    "coverage": 0.0,
                    "paragraphs": [{
                        "text": f"[MISSING] Provide evidence for {section_key}",
                        "citations": []
                    }]
                }
        
        return sections
    
    def _get_section_keys(self, doc_type: str) -> List[str]:
        """Get section keys for a document type."""
        section_mapping = {
            "annex_iv": [
                "section_2_1_architecture", "section_2_2_data_processing",
                "section_2_3_model_info", "section_3_1_risk_assessment",
                "section_3_2_mitigation", "section_4_1_data_sources",
                "section_4_2_data_quality", "section_4_3_data_retention",
                "section_5_1_decision_logic", "section_5_2_explainability",
                "section_6_1_human_oversight", "section_6_2_override",
                "section_7_1_testing", "section_7_2_validation",
                "section_8_1_monitoring", "section_8_2_logging",
                "section_8_3_supervision", "section_9_1_iso_compliance"
            ],
            "fria": [
                "section_1_1_system_description", "section_1_2_intended_use",
                "section_1_3_target_users", "section_2_1_privacy_rights",
                "section_2_2_non_discrimination", "section_2_3_freedom_expression",
                "section_2_4_right_information", "section_3_1_identified_risks",
                "section_3_2_risk_severity", "section_3_3_risk_likelihood",
                "section_4_1_affected_groups", "section_4_2_impact_magnitude",
                "section_4_3_duration_impact", "section_5_1_technical_safeguards",
                "section_5_2_organizational_measures", "section_5_3_legal_protections",
                "section_6_1_monitoring_plan", "section_6_2_review_schedule",
                "section_6_3_update_procedures", "section_7_1_consultation_process",
                "section_7_2_stakeholder_feedback", "section_7_3_response_concerns",
                "section_8_1_overall_assessment", "section_8_2_recommendations",
                "section_8_3_implementation_timeline"
            ],
            "pmm": [
                "section_1_1_operational_metrics", "section_1_2_performance_trends",
                "section_1_3_system_reliability", "section_2_1_incident_summary",
                "section_2_2_incident_classification", "section_2_3_root_cause_analysis",
                "section_2_4_corrective_actions", "section_3_1_risk_indicators",
                "section_3_2_risk_trends", "section_3_3_emerging_risks",
                "section_4_1_regulatory_compliance", "section_4_2_standards_adherence",
                "section_4_3_audit_findings", "section_5_1_user_feedback",
                "section_5_2_complaints_analysis", "section_5_3_satisfaction_metrics",
                "section_6_1_data_quality_metrics", "section_6_2_data_drift_detection",
                "section_6_3_model_performance", "section_7_1_planned_improvements",
                "section_7_2_implementation_status", "section_7_3_effectiveness_review",
                "section_8_1_internal_reporting", "section_8_2_external_communication",
                "section_8_3_regulatory_notifications"
            ],
            "soa": [
                "section_1_1_organization_context", "section_1_2_ai_systems_inventory",
                "section_1_3_stakeholder_analysis", "section_2_1_leadership_commitment",
                "section_2_2_policy", "section_2_3_roles_responsibilities",
                "section_3_1_risk_assessment", "section_3_2_risk_treatment",
                "section_3_3_risk_treatment_plan", "section_4_1_resources",
                "section_4_2_competence", "section_4_3_awareness",
                "section_4_4_communication", "section_4_5_documented_information",
                "section_5_1_operational_planning", "section_5_2_impact_assessment",
                "section_5_3_risk_treatment_implementation", "section_5_4_development_lifecycle",
                "section_5_5_ai_system_use", "section_6_1_monitoring_measurement",
                "section_6_2_internal_audit", "section_6_3_management_review",
                "section_7_1_nonconformity_corrective", "section_7_2_continual_improvement",
                "section_8_1_high_priority_controls", "section_8_2_medium_priority_controls",
                "section_8_3_low_priority_controls", "section_9_1_completed_controls",
                "section_9_2_in_progress_controls", "section_9_3_planned_controls"
            ],
            "risk_register": [
                "section_1_1_high_risk_items", "section_1_2_medium_risk_items",
                "section_1_3_low_risk_items", "section_2_1_risk_identification",
                "section_2_2_risk_analysis", "section_2_3_risk_evaluation",
                "section_3_1_treatment_options", "section_3_2_treatment_plans",
                "section_3_3_treatment_implementation", "section_4_1_capa_process",
                "section_4_2_open_capas", "section_4_3_closed_capas",
                "section_5_1_risk_monitoring_plan", "section_5_2_risk_review_schedule",
                "section_5_3_risk_escalation_procedures", "section_6_1_incident_classification",
                "section_6_2_incident_response", "section_6_3_incident_investigation",
                "section_7_1_internal_communication", "section_7_2_external_communication",
                "section_7_3_stakeholder_reporting", "section_8_1_risk_culture_assessment",
                "section_8_2_training_awareness", "section_8_3_risk_governance"
            ]
        }
        
        return section_mapping.get(doc_type, [])
    
    def _search_evidence_snippets(
        self, 
        db: Session, 
        org_id: int, 
        system_id: Optional[int], 
        section_key: str
    ) -> List[Dict]:
        """Search for evidence snippets matching the section key using ArtifactText."""
        from app.services.text_extraction import search_artifact_text
        import re
        
        # Extract meaningful keywords from section_key
        # e.g., "section_2_1_architecture" → "architecture"
        # Remove "section_" prefix and numeric parts
        keywords = section_key.replace("section_", "")
        keywords = re.sub(r'\d+_\d+_', '', keywords)  # Remove pattern like "2_1_"
        keywords = keywords.replace("_", " ")
        
        # Search ArtifactText for relevant content
        artifacts = search_artifact_text(
            db=db,
            org_id=org_id,
            system_id=system_id,
            search_term=keywords,
            limit=3
        )
        
        # Convert to snippets format
        return [
            {
                "evidence_id": artifact.evidence_id,
                "page_number": artifact.page,
                "text_content": artifact.content[:500],  # Truncate to 500 chars
                "checksum": artifact.checksum
            }
            for artifact in artifacts
        ]
    
    def _generate_evidence_paragraph(
        self, 
        evidence_snippets: List[Dict], 
        section_key: str
    ) -> Dict:
        """Generate an evidence-grounded paragraph with citations."""
        
        if not evidence_snippets:
            return {
                "text": f"[MISSING] Provide evidence for {section_key}",
                "citations": []
            }
        
        # Combine text snippets
        combined_text = " ".join([
            snippet["text_content"] for snippet in evidence_snippets
        ])
        
        # Generate citations
        citations = [
            {
                "evidence_id": snippet["evidence_id"],
                "page": snippet["page_number"],
                "checksum": snippet["checksum"]
            }
            for snippet in evidence_snippets
        ]
        
        # Format paragraph with citations
        citation_refs = []
        for i, citation in enumerate(citations, 1):
            ref = f"[EV-{citation['evidence_id']} p.{citation['page']} | sha256:{citation['checksum'][:16]}]"
            citation_refs.append(ref)
        
        # Append citations to text
        text_with_citations = f"{combined_text} {' '.join(citation_refs)}"
        
        return {
            "text": text_with_citations,
            "citations": citations
        }
    
    def _prepare_template_variables(
        self, 
        org: Organization, 
        system: Optional[AISystem], 
        sections: Dict[str, Dict]
    ) -> Dict:
        """Prepare template variables for document generation."""
        
        now = datetime.now(timezone.utc)
        
        # Base variables
        vars = {
            "org_name": org.name,
            "generated_at": now.isoformat(),
            "timestamp_utc": now.isoformat(),
            "bundle_hash": self._generate_bundle_hash(sections)
        }
        
        # System-specific variables
        if system:
            vars.update({
                "system_name": system.name,
                "system_purpose": system.purpose,
                "ai_act_class": getattr(system, 'ai_act_class', 'Not assessed'),
                "is_gpai": getattr(system, 'is_gpai', False),
                "role": getattr(system, 'role', 'Not determined')
            })
        
        # Add section content
        for section_key, section_data in sections.items():
            if section_data.get("paragraphs"):
                vars[section_key] = section_data["paragraphs"][0]["text"]
            else:
                vars[section_key] = f"[MISSING] Provide evidence for {section_key}"
        
        # Add placeholder content for missing sections
        for section_key, section_data in sections.items():
            if section_data.get("coverage", 0) == 0:
                vars[section_key] = f"[MISSING] Provide evidence for {section_key}"
        
        return vars
    
    def _generate_bundle_hash(self, sections: Dict[str, Dict]) -> str:
        """Generate SHA-256 hash of the document bundle."""
        content = json.dumps(sections, sort_keys=True)
        return hashlib.sha256(content.encode()).hexdigest()[:16]
    
    def export_document(
        self, 
        db: Session, 
        org_id: int, 
        system_id: Optional[int], 
        doc_type: str, 
        format: str
    ) -> Tuple[str, bytes]:
        """
        Export a document in the specified format.
        
        Args:
            db: Database session
            org_id: Organization ID
            system_id: Optional system ID
            doc_type: Document type (annex_iv, fria, pmm, soa, risk_register)
            format: Export format (md, docx, pdf)
            
        Returns:
            Tuple of (filename, content_bytes)
        """
        
        # Validate system ownership if system_id provided
        system = None
        if system_id:
            system = db.query(AISystem).filter(
                AISystem.id == system_id,
                AISystem.org_id == org_id  # Enforce org scoping to prevent cross-tenant leak
            ).first()
            if not system:
                raise ValueError(f"System {system_id} not found or access denied")
        
        # Generate document content
        doc_result = self._generate_document(
            db, 
            db.query(Organization).filter(Organization.id == org_id).first(),
            system,  # Use validated system
            doc_type
        )
        
        content = doc_result["content"]
        
        # Generate filename
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
        system_suffix = f"-{system_id}" if system_id else ""
        filename = f"{doc_type}{system_suffix}-{timestamp}.{format}"
        
        # Convert to requested format
        if format == "md":
            return filename, content.encode("utf-8")
        elif format == "docx":
            return self._convert_to_docx(content, filename)
        elif format == "pdf":
            return self._convert_to_pdf(content, filename)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _convert_to_docx(self, content: str, filename: str) -> Tuple[str, bytes]:
        """Convert markdown content to DOCX format."""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Split content into lines
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                else:
                    doc.add_paragraph(line)
            
            # Save to bytes
            import io
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            
            return filename, bio.getvalue()
            
        except ImportError:
            raise ValueError("python-docx not available for DOCX export")
    
    def _convert_to_pdf(self, content: str, filename: str) -> Tuple[str, bytes]:
        """Convert markdown content to PDF format."""
        try:
            import markdown
            from weasyprint import HTML, CSS
            
            if not settings.ENABLE_PDF_EXPORT:
                raise ValueError("PDF export is disabled")
            
            # Convert markdown to HTML
            html_content = markdown.markdown(content)
            
            # Add CSS styling
            css = CSS(string='''
                body { font-family: Arial, sans-serif; margin: 40px; }
                h1 { color: #333; border-bottom: 2px solid #333; }
                h2 { color: #666; }
                h3 { color: #999; }
                p { line-height: 1.6; }
                code { background-color: #f5f5f5; padding: 2px 4px; }
            ''')
            
            # Generate PDF
            html_doc = HTML(string=html_content)
            pdf_bytes = html_doc.write_pdf(stylesheets=[css])
            
            return filename, pdf_bytes
            
        except ImportError:
            raise ValueError("WeasyPrint not available for PDF export")


# Global service instance
compliance_suite_service = ComplianceSuiteService()
